# Импорт необходимых библиотек
from dotenv import load_dotenv
load_dotenv()

import logging  # Для логирования работы бота
import os  # Для работы с переменными окружения и файловой системой
import base64
import sqlite3  # Для работы с SQLite базой данных
from datetime import datetime  # Для работы с датой и временем

import pandas as pd  # Для обработки данных и экспорта в Excel
from docx import Document  # Для создания Word-документов
from github import Github, UnknownObjectException, GithubException  # Для взаимодействия с GitHub API
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup  # Компоненты Telegram Bot API
from telegram.ext import (
    Application,
    ContextTypes,
    CommandHandler,
    MessageHandler,
    filters,
    ConversationHandler,
    CallbackQueryHandler
)  # Обработчики команд и сообщений

# Конфиdoгурационные параметры (берутся из переменных окружения)
TOKEN = os.getenv('TELEGRAM_TOKEN')  # Токен бота
ADMIN_IDS = [429442647]  # Список ID администраторов
GITHUB_TOKEN = os.getenv('GITHUB_TOKEN')  # Токен для доступа к GitHub
REPO_NAME = os.getenv('GITHUB_REPO_URL')  # Название репозитория GitHub
DB_FILE = "repair_requests.db"  # Название файла базы данных

# Определение состояний для ConversationHandler
CONTACT, DESCRIPTION, TIME = range(3)  # Этапы диалога: контакты, описание, время

# Настройка логирования
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)


class DatabaseManager:
    """Класс для управления базой данных SQLite"""

    def __init__(self, db_name):
        # Регистрация кастомных адаптеров для работы с datetime
        sqlite3.register_adapter(datetime, lambda val: val.isoformat())
        sqlite3.register_converter("datetime", lambda b: datetime.fromisoformat(b.decode()))

        # Подключение к базе данных с поддержкой типов
        self.conn = sqlite3.connect(db_name, detect_types=sqlite3.PARSE_DECLTYPES)
        self._create_table()

    def _create_table(self):
        """Создание таблицы запросов, если она не существует"""
        cursor = self.conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS requests (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                username TEXT,
                contact_info TEXT,
                problem_description TEXT,
                preferred_time TEXT,
                created_at DATETIME
            )
        ''')
        self.conn.commit()

    def add_request(self, user_data):
        """Добавление новой заявки в базу данных"""
        cursor = self.conn.cursor()
        cursor.execute('''
            INSERT INTO requests 
            (user_id, username, contact_info, problem_description, preferred_time, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (
            user_data['user_id'],
            user_data['username'],
            user_data['contact_info'],
            user_data['description'],
            user_data['time'],
            datetime.now()
        ))
        self.conn.commit()

    def get_all_requests(self):
        """Получение всех заявок из базы данных"""
        cursor = self.conn.cursor()
        cursor.execute('SELECT * FROM requests')
        return cursor.fetchall()


# Инициализация объекта базы данных
db = DatabaseManager(DB_FILE)


# Клавиатура для отмены действий.
cancel_keyboard = InlineKeyboardMarkup([
    [InlineKeyboardButton("❌ Отмена", callback_data='cancel')]
])


### Обработчики команд ###
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /start"""
    user = update.effective_user
    text = (
        f"Привет, {user.first_name}! Я бот для оформления заявок на ремонт.\n"
        "Доступные команды:\n"
        "/new_request - создать новую заявку\n"
        "/help - справка по использованию"
    )

    # Добавляем команду экспорта для админов
    if user.id in ADMIN_IDS:
        text += "\n/export - экспорт данных (только для администраторов)"

    await update.message.reply_text(text)

async def show_help(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработчик команды /help"""
    help_text = (
        "🔧 Этот бот помогает создавать заявки на ремонт бытовой техники.\n\n"
        "Основные команды:\n"
        "/start - начать работу с ботом\n"
        "/new_request - создать новую заявку\n"
        "/help - показать эту справку\n\n"
        "Вы также можете использовать текстовые команды:\n"
        "'ремонт' - начать новую заявку\n"
        "'помощь' - показать справку"
    )
    await update.message.reply_text(help_text)

async def new_request(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Начало создания новой заявки"""
    await update.message.reply_text(
        "Пожалуйста, введите ваши контактные данные (телефон или email):",
        reply_markup=cancel_keyboard
    )
    return CONTACT

async def contact(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка контактных данных"""
    contact_info = update.message.text
    # Простая валидация контактов
    if not ('@' in contact_info or any(c.isdigit() for c in contact_info)):
        await update.message.reply_text("Некорректные данные. Повторите:", reply_markup=cancel_keyboard)
        return CONTACT
    context.user_data['contact_info'] = contact_info
    await update.message.reply_text("Опишите проблему:", reply_markup=cancel_keyboard)
    return DESCRIPTION

async def description(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка описания проблемы"""
    if update.callback_query and update.callback_query.data == 'cancel':
        return await cancel(update, context)

    context.user_data['description'] = update.message.text
    await update.message.reply_text(
        "Укажите предпочтительное время для связи:",
        reply_markup=cancel_keyboard
    )
    return TIME

async def time(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка предпочтительного времени"""
    if update.callback_query and update.callback_query.data == 'cancel':
        return await cancel(update, context)

    context.user_data['time'] = update.message.text
    user_data = {
        'user_id': update.effective_user.id,
        'username': update.effective_user.username,
        **context.user_data
    }

    db.add_request(user_data)
    await update.message.reply_text("✅ Заявка успешно создана!")
    return ConversationHandler.END

async def export_data(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Экспорт данных (доступно только администраторам)"""
    if update.effective_user.id not in ADMIN_IDS:
        await update.message.reply_text("⛔ У вас нет прав для выполнения этой команды")
        return

    await update.message.reply_text(
        "Выберите формат экспорта:",
        reply_markup=InlineKeyboardMarkup([
            [
                InlineKeyboardButton("Excel", callback_data='export_excel'),
                InlineKeyboardButton("Word", callback_data='export_doc')
            ],
            [InlineKeyboardButton("Text", callback_data='export_txt')]
        ])
    )

async def export_callback(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Обработка выбора формата экспорта"""
    global filename, file_type
    query = update.callback_query
    await query.answer()

    # Проверка допустимых форматов
    if query.data not in ['export_excel', 'export_doc', 'export_txt']:
        await query.message.edit_text("Неизвестная команда")
        return

    # Получение данных из БД
    requests = db.get_all_requests()
    df = pd.DataFrame(requests, columns=[
        'ID', 'User ID', 'Username', 'Contact', 'Description',
        'Preferred Time', 'Created At'
    ])

    try:
        # Генерация файла в выбранном формате
        if query.data == 'export_excel':
            filename = f"requests_{datetime.now().strftime('%Y%m%d%H%M')}.xlsx"
            df.to_excel(filename, index=False)
            file_type = "Excel"

        elif query.data == 'export_doc':
            filename = f"requests_{datetime.now().strftime('%Y%m%d%H%M')}.docx"
            doc = Document()
            doc.add_heading('Заявки на ремонт', 0)
            for _, row in df.iterrows():
                doc.add_paragraph(
                    f"ID: {row['ID']}\n"
                    f"User ID: {row['User ID']}\n"
                    f"Username: {row['Username']}\n"
                    f"Contact: {row['Contact']}\n"
                    f"Description: {row['Description']}\n"
                    f"Preferred Time: {row['Preferred Time']}\n"
                    f"Created At: {row['Created At']}\n"
                    "\n---\n"
                )
            doc.save(filename)
            file_type = "Word"

        elif query.data == 'export_txt':
            filename = f"requests_{datetime.now().strftime('%Y%m%d%H%M')}.txt"
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(df.to_string(index=False))
            file_type = "Text"

        # Отправка файла пользователю
        with open(filename, 'rb') as f:
            await query.message.reply_document(
                document=f,
                caption=f"Экспорт в формате {file_type} выполнен ✅"
            )

        # Удаление кнопок после выбора
        await query.message.edit_reply_markup(reply_markup=None)

    except Exception as e:
        logging.error(f"Export error: {str(e)}")
        await query.message.reply_text("❌ Ошибка при экспорте")

    finally:
        # Удаление временного файла
        if 'filename' in locals():
            if os.path.exists(filename):
                os.remove(filename)


async def sync_with_github(context: ContextTypes.DEFAULT_TYPE):
    """Синхронизация базы данных с GitHub"""
    try:
        # Проверка наличия обязательных переменных
        if not GITHUB_TOKEN or not REPO_NAME:
            raise ValueError("Missing GitHub credentials in environment variables")

        # Инициализация подключения
        g = Github(GITHUB_TOKEN)

        try:
            repo = g.get_repo(REPO_NAME)
        except UnknownObjectException:
            raise ValueError(f"Repository {REPO_NAME} not found")

        # Проверка существования локального файла
        if not os.path.exists(DB_FILE):
            raise FileNotFoundError(f"Local database file {DB_FILE} not found")

        # Чтение и кодирование содержимого файла
        with open(DB_FILE, 'rb') as f:
            content = f.read()
        encoded_content = base64.b64encode(content).decode('utf-8')

        # Попытка обновления существующего файла
        try:
            contents = repo.get_contents(DB_FILE, ref="main")
            update_result = repo.update_file(
                path=DB_FILE,
                message="Automatic DB sync",
                content=encoded_content,
                sha=contents.sha,
                branch="main"
            )
            logging.info(f"File updated: {update_result['commit'].html_url}")

        except UnknownObjectException:
            # Создание нового файла если не существует
            create_result = repo.create_file(
                path=DB_FILE,
                message="Automatic DB sync",
                content=encoded_content,
                branch="main"
            )
            logging.info(f"File created: {create_result['commit'].html_url}")

    except GithubException as ge:
        logging.error(f"GitHub API Error: {ge.data.get('message')}")
        # Можно добавить уведомление админу через Telegram
        await context.bot.send_message(
            chat_id=ADMIN_IDS[0],
            text=f"⚠️ GitHub sync failed: {ge.data.get('message')}"
        )

    except FileNotFoundError as fe:
        logging.error(f"Local file error: {str(fe)}")

    except ValueError as ve:
        logging.error(f"Configuration error: {str(ve)}")

    except Exception as e:
        logging.error(f"Unexpected error: {str(e)}")
        await context.bot.send_message(
            chat_id=ADMIN_IDS[0],
            text=f"⚠️ Critical sync error: {str(e)}"
        )

    else:
        logging.info("Database sync completed successfully")


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Отмена текущего действия"""
    query = update.callback_query
    if query:
        await query.answer()
        await query.message.edit_reply_markup(reply_markup=None)
        await query.message.reply_text("🚫 Создание заявки отменено")
    else:
        await update.message.reply_text("🚫 Создание заявки отменено")

    context.user_data.clear()
    return ConversationHandler.END


def main():
    """Основная функция инициализации бота"""
    # Создание приложения бота
    application = Application.builder().token(TOKEN).build()

    # Настройка ConversationHandler для диалога создания заявки
    conv_handler = ConversationHandler(
        entry_points=[
            CommandHandler('new_request', new_request),
            MessageHandler(filters.Regex(r'(?i)\b(ремонт[ау]?|сломал(ся|ась)?|поломк[аи])\b|не\s+работает'), new_request)
        ],
        states={
            CONTACT: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, contact),
                CallbackQueryHandler(cancel, pattern='^cancel$')
            ],
            DESCRIPTION: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, description),
                CallbackQueryHandler(cancel, pattern='^cancel$')
            ],
            TIME: [
                MessageHandler(filters.TEXT & ~filters.COMMAND, time),
                CallbackQueryHandler(cancel, pattern='^cancel$')
            ]
        },
        fallbacks=[CommandHandler('cancel', cancel)]
    )

    # Регистрация обработчиков команд
    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', show_help))
    application.add_handler(CommandHandler('export', export_data))
    application.add_handler(CallbackQueryHandler(export_callback, pattern="^export_"))
    application.add_handler(conv_handler)

    # Настройка периодической синхронизации с GitHub (каждый час)
    application.job_queue.run_repeating(
        sync_with_github,
        interval=3600,
        first=10
    )

    # Запуск бота
    application.run_polling()

if __name__ == '__main__':
    main()